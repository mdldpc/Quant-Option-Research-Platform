# ============================================
# Seven Goal Report Builder v1.0
# ============================================

from pathlib import Path
from docx import Document
from docx.shared import Inches


# --------------------------------------------
# Paths
# --------------------------------------------

PROJECT_ROOT = Path(__file__).resolve().parents[1]

INPUT_MD = PROJECT_ROOT / "research" / "notebooks" / "Seven_Goal_Project_Report.md"

OUTPUT_DOCX = (
    PROJECT_ROOT
    / "research"
    / "reports"
    / "Seven_Goal_Project_Report_v1.docx"
)


# --------------------------------------------
# Load Markdown
# --------------------------------------------

def load_markdown(path):

    with open(path, "r", encoding="utf-8") as f:
        return f.readlines()


# --------------------------------------------
# Image parser
# --------------------------------------------

def parse_image(line):

    if "![" not in line:
        return None

    left = line.find("(")
    right = line.find(")")

    if left == -1 or right == -1:
        return None

    return line[left + 1:right]


# --------------------------------------------
# Build Document
# --------------------------------------------

def build_document(lines):

    doc = Document()

    for line in lines:

        line = line.rstrip()

        if line == "":
            continue

        # -----------------------
        # Heading 1
        # -----------------------

        if line.startswith("# "):

            doc.add_heading(line[2:], level=1)
            continue

        # -----------------------
        # Heading 2
        # -----------------------

        if line.startswith("## "):

            doc.add_heading(line[3:], level=2)
            continue

        # -----------------------
        # Heading 3
        # -----------------------

        if line.startswith("### "):

            doc.add_heading(line[4:], level=3)
            continue

        # -----------------------
        # Image
        # -----------------------

        image = parse_image(line)

        if image:

            image_path = INPUT_MD.parent / image

            if image_path.exists():

                doc.add_picture(
                    str(image_path),
                    width=Inches(6.5)
                )

            else:

                doc.add_paragraph(
                    f"[Missing image: {image}]"
                )

            continue

        # -----------------------
        # Normal paragraph
        # -----------------------

        doc.add_paragraph(line)

    return doc


# --------------------------------------------
# Main
# --------------------------------------------

def main():

    print("Reading markdown...")

    lines = load_markdown(INPUT_MD)

    print("Building Word document...")

    doc = build_document(lines)

    OUTPUT_DOCX.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    doc.save(OUTPUT_DOCX)

    print()

    print("Saved:")

    print(OUTPUT_DOCX)

    print()

    print("DONE")


if __name__ == "__main__":

    main()