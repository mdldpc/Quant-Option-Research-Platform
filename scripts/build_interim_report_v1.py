from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.paths import PROJECT_ROOT


OUTPUT_DIR = PROJECT_ROOT / "research" / "final_reports" / "interim_report_v1"
OUTPUT_DOCX = OUTPUT_DIR / "Interim_Research_Report_v1.docx"

RQ1_PATH = PROJECT_ROOT / "research" / "notebooks" / "RQ1_Daily_ATM_IV_Distribution.md"

FIGURE_1 = PROJECT_ROOT / "research" / "figures" / "iv_distribution_histogram.png"
FIGURE_2 = PROJECT_ROOT / "research" / "figures" / "iv_distribution_boxplot.png"


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quant Option Research Platform")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Interim Research Report v1.0")
    run.font.size = Pt(16)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Jingzhe Yang\n")
    meta.add_run("Version: Interim Report v1.0\n")
    meta.add_run("Status: Research Draft\n")

    doc.add_page_break()


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    if text.strip():
        doc.add_paragraph(text.strip())


def add_figure(doc, image_path, caption):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))

        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Missing figure: {image_path}]")


def add_rq1_from_markdown(doc):
    text = RQ1_PATH.read_text(encoding="utf-8")

    add_heading(doc, "Research Question 1: Daily ATM Implied Volatility Distribution", 1)

    for line in text.splitlines():
        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            add_heading(doc, line.replace("# ", ""), 1)

        elif line.startswith("## "):
            add_heading(doc, line.replace("## ", ""), 2)

        elif line.startswith("### "):
            add_heading(doc, line.replace("### ", ""), 3)

        elif line.startswith("!["):
            # Skip markdown image syntax; figures are inserted manually below.
            continue

        elif line.startswith("*Figure"):
            # Skip markdown figure captions; figures are inserted manually below.
            continue

        elif line.startswith("**Figure 5.1"):
            add_figure(
                doc,
                FIGURE_1,
                "Figure 5.1. Distribution of Daily ATM Implied Volatility",
            )

        elif line.startswith("**Figure 5.2"):
            add_figure(
                doc,
                FIGURE_2,
                "Figure 5.2. Box Plot of Daily ATM Implied Volatility",
            )

        elif line.startswith("|"):
            # Skip markdown tables in v1; we will improve table rendering later.
            continue

        else:
            add_paragraph(doc, line)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    add_title_page(doc)

    add_heading(doc, "Executive Summary", 1)
    add_paragraph(
        doc,
        "This interim report summarizes the first completed research release "
        "of the Quant Option Research Platform. The report combines the "
        "engineering platform developed in Version 1.0 with the first formal "
        "research question completed in Version 2.0."
    )

    add_heading(doc, "Project Overview", 1)
    add_paragraph(
        doc,
        "The Quant Option Research Platform is designed for option volatility "
        "modelling, Greeks calculation, volatility surface construction, "
        "strategy backtesting, robustness analysis, and systematic research "
        "reporting."
    )

    add_heading(doc, "Current Research Scope", 1)
    add_paragraph(
        doc,
        "The current report focuses on Research Question 1, which investigates "
        "the empirical distribution of daily at-the-money implied volatility "
        "using the available 2026 Chinese index option dataset."
    )

    add_rq1_from_markdown(doc)

    doc.save(OUTPUT_DOCX)

    print("DONE")
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()