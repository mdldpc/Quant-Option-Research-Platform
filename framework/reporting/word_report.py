"""
Word Report Builder

Reusable helper for generating Word (.docx) reports.

This module intentionally contains no business logic.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


class WordReport:

    def __init__(self, language="en"):
        """
        Parameters
        ----------
        language : str
            "en" for English reports.
            "zh" for Chinese reports.
        """
        self.language = language.lower()
        self.document = Document()

        if self.language == "zh":
            self.western_font = "Calibri"
            self.east_asia_font = "Microsoft YaHei"
        else:
            self.western_font = "Calibri"
            self.east_asia_font = "Calibri"

        self._configure_styles()

    # --------------------------------------------------
    # Font helpers
    # --------------------------------------------------

    def _configure_style(
        self,
        style_name,
        size=None,
        bold=None,
    ):
        """
        Configure both Western and East Asian fonts for a Word style.
        """
        try:
            style = self.document.styles[style_name]
        except KeyError:
            return

        style.font.name = self.western_font

        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()

        rfonts.set(
            qn("w:ascii"),
            self.western_font,
        )
        rfonts.set(
            qn("w:hAnsi"),
            self.western_font,
        )
        rfonts.set(
            qn("w:eastAsia"),
            self.east_asia_font,
        )

        if size is not None:
            style.font.size = Pt(size)

        if bold is not None:
            style.font.bold = bold

    def _configure_styles(self):
        """
        Configure fonts for all commonly used document styles.
        """
        self._configure_style(
            "Normal",
            size=11,
        )

        self._configure_style(
            "Title",
            size=24,
            bold=True,
        )

        self._configure_style(
            "Subtitle",
            size=14,
        )

        self._configure_style(
            "Heading 1",
            size=16,
            bold=True,
        )

        self._configure_style(
            "Heading 2",
            size=13,
            bold=True,
        )

        self._configure_style(
            "Heading 3",
            size=11,
            bold=True,
        )

        self._configure_style(
            "List Bullet",
            size=11,
        )

        self._configure_style(
            "List Number",
            size=11,
        )

        self._configure_style(
            "Caption",
            size=9,
        )

        self._configure_style(
            "Table Grid",
            size=10,
        )

    def _set_run_font(
        self,
        run,
        size=None,
        bold=None,
        italic=None,
        font_name=None,
        east_asia_font=None,
    ):
        """
        Apply Western and East Asian font settings directly to a run.
        """
        western_font = font_name or self.western_font
        chinese_font = east_asia_font or self.east_asia_font

        run.font.name = western_font

        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()

        rfonts.set(
            qn("w:ascii"),
            western_font,
        )
        rfonts.set(
            qn("w:hAnsi"),
            western_font,
        )
        rfonts.set(
            qn("w:eastAsia"),
            chinese_font,
        )

        if size is not None:
            run.font.size = Pt(size)

        if bold is not None:
            run.bold = bold

        if italic is not None:
            run.italic = italic

        return run

    def _format_cell_runs(
        self,
        cell,
        bold=False,
        size=9,
    ):
        """
        Ensure all text inside a table cell uses the correct bilingual font.
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                self._set_run_font(
                    run,
                    size=size,
                    bold=bold,
                )

    # --------------------------------------------------
    # Headings
    # --------------------------------------------------

    def title(self, text):
        p = self.document.add_heading(level=0)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = p.add_run(str(text))

        self._set_run_font(
            run,
            size=24,
            bold=True,
        )

        return p

    def heading1(self, text):
        p = self.document.add_heading(
            str(text),
            level=1,
        )

        for run in p.runs:
            self._set_run_font(
                run,
                size=16,
                bold=True,
            )

        return p

    def heading2(self, text):
        p = self.document.add_heading(
            str(text),
            level=2,
        )

        for run in p.runs:
            self._set_run_font(
                run,
                size=13,
                bold=True,
            )

        return p

    def heading3(self, text):
        p = self.document.add_heading(
            str(text),
            level=3,
        )

        for run in p.runs:
            self._set_run_font(
                run,
                size=11,
                bold=True,
            )

        return p

    # --------------------------------------------------
    # Paragraph
    # --------------------------------------------------

    def paragraph(
        self,
        text="",
        bold=False,
        italic=False,
        alignment=None,
    ):
        p = self.document.add_paragraph()

        if alignment is not None:
            p.alignment = alignment

        run = p.add_run(str(text))

        self._set_run_font(
            run,
            size=11,
            bold=bold,
            italic=italic,
        )

        return p

    def code_block(self, text):
        p = self.document.add_paragraph()
        run = p.add_run(str(text))

        self._set_run_font(
            run,
            size=9,
            font_name="Consolas",
            east_asia_font=self.east_asia_font,
        )

        return p

    # --------------------------------------------------
    # Bullet / Numbered
    # --------------------------------------------------

    def bullet(self, text):
        p = self.document.add_paragraph(
            style="List Bullet",
        )

        run = p.add_run(str(text))

        self._set_run_font(
            run,
            size=11,
        )

        return p

    def numbered(self, text):
        p = self.document.add_paragraph(
            style="List Number",
        )

        run = p.add_run(str(text))

        self._set_run_font(
            run,
            size=11,
        )

        return p

    # --------------------------------------------------
    # Basic Table
    # --------------------------------------------------

    def dataframe(self, df):
        """
        Backward-compatible DataFrame renderer.
        """
        return self._add_dataframe(df)

    def _add_dataframe(self, df):
        if df is None:
            return None

        if len(df.columns) == 0:
            return None

        table = self.document.add_table(
            rows=1,
            cols=len(df.columns),
        )

        table.style = "Table Grid"

        header_cells = table.rows[0].cells

        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            header_cells[i].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            self._format_cell_runs(
                header_cells[i],
                bold=True,
                size=9,
            )

        for _, row in df.iterrows():
            cells = table.add_row().cells

            for i, value in enumerate(row):
                cells[i].text = (
                    ""
                    if value is None
                    else str(value)
                )

                cells[i].vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.CENTER
                )

                self._format_cell_runs(
                    cells[i],
                    bold=False,
                    size=9,
                )

        return table

    # --------------------------------------------------
    # Publication Table
    # --------------------------------------------------

    def table(
        self,
        dataframe,
        title=None,
        caption=None,
    ):
        if title:
            title_paragraph = self.document.add_paragraph()
            title_paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            title_run = title_paragraph.add_run(str(title))

            self._set_run_font(
                title_run,
                size=10,
                bold=True,
            )

        table = self._add_dataframe(dataframe)

        if caption:
            caption_paragraph = self.document.add_paragraph()
            caption_paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            caption_run = caption_paragraph.add_run(
                str(caption)
            )

            self._set_run_font(
                caption_run,
                size=9,
                italic=True,
            )

        return table

    # --------------------------------------------------
    # Basic Image
    # --------------------------------------------------

    def image(
        self,
        image_path,
        width=6.5,
    ):
        image_path = Path(image_path)

        if not image_path.exists():
            return None

        paragraph = self.document.add_paragraph()
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run = paragraph.add_run()

        return run.add_picture(
            str(image_path),
            width=Inches(width),
        )

    # --------------------------------------------------
    # Publication Figure
    # --------------------------------------------------

    def figure(
        self,
        image,
        title=None,
        caption=None,
        width=6.5,
    ):
        image_path = Path(image)

        if not image_path.exists():
            return None

        if title:
            title_paragraph = self.document.add_paragraph()
            title_paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            title_run = title_paragraph.add_run(
                str(title)
            )

            self._set_run_font(
                title_run,
                size=10,
                bold=True,
            )

        image_paragraph = self.document.add_paragraph()
        image_paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        image_run = image_paragraph.add_run()

        picture = image_run.add_picture(
            str(image_path),
            width=Inches(width),
        )

        if caption:
            caption_paragraph = self.document.add_paragraph()
            caption_paragraph.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )

            caption_run = caption_paragraph.add_run(
                str(caption)
            )

            self._set_run_font(
                caption_run,
                size=9,
                italic=True,
            )

        return picture

    # --------------------------------------------------
    # Misc
    # --------------------------------------------------

    def page_break(self):
        self.document.add_page_break()

    def horizontal_rule(self):
        self.paragraph("-" * 80)

    # --------------------------------------------------
    # Save
    # --------------------------------------------------

    def save(self, path):
        path = Path(path)

        path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        self.document.save(path)